"""SAFE 개선 로드맵 — Word(.docx) 생성 스크립트"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent.parent / "docs"
OUT_NAMES = ["SAFE_개선로드맵.docx", "SAFE_improvement_roadmap.docx"]


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build() -> Path:
    doc = Document()
    title = doc.add_heading("SAFE 대규모 개선 로드맵", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = meta.add_run(
        "작성 목적: 지방보조사업 검토 테스트 보완 사항의 단계별 추진 기준\n"
        "작성일: 2026-06-29  |  상태: 기획·요구정리 단계 (코드 미반영)\n"
        "관련 문서: manual.md (현행 SAFE 사용 설명서)"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_heading(doc, "1. 배경 — 왜 대대적 수정이 필요한가", 1)
    add_para(
        doc,
        "몇 차례 지방보조사업 검토 테스트 결과, 다음 한계가 기능 몇 개 추가로는 "
        "해결되지 않음을 확인했습니다.",
    )
    add_table(
        doc,
        ["#", "보완점", "요약"],
        [
            [
                "1",
                "SILO 형태의 검토",
                "4종 자료를 각각 따로 검토·결과 관리. 실무는 사업 1건의 4종 자료를 "
                "종합하여 적정/부적정 판단",
            ],
            [
                "2",
                "사업별 특성 미고려",
                "비목·사업 유형과 무관하게 동일 체크리스트·규칙 적용 → 해당 없는 항목까지 "
                "부적정 판정 (예: 공사비만 집행하는 사업에 원천징수 미징수 부적정)",
            ],
        ],
    )
    add_para(doc, "두 문제의 공통 뿌리:", bold=True)
    add_bullets(
        doc,
        [
            "「무엇을 검토하는가」→ 자료 1건이 아니라 사업(케이스) 1건",
            "「무엇을 기준으로 판단하는가」→ 일괄 체크리스트가 아니라 "
            "사업 프로필·적용 조건이 반영된 점검",
        ],
    )

    add_heading(doc, "2. 현행 SAFE 구조 (개선 전 기준선)", 1)
    add_table(
        doc,
        ["항목", "현행"],
        [
            ["검토 단위", "자료유형 1개 + 파일 1건 (③만 ZIP 묶음 가능)"],
            ["체크리스트", "유형별(①~④) DB 저장, 검토 시 선택"],
            ["판정", "체크리스트 항목 유사도 + 코드 고정 규칙검증"],
            ["지식DB", "체크리스트 생성 시 참조. 검토 비교 단계에서는 미사용"],
            ["Gemini", "검토 파이프라인 미포함. 결과 화면 추가분석(항목명만)만 선택적"],
            ["교차 검토", "없음 (계획서 vs 집행내역 등 대조 불가)"],
            ["규칙검증", "checker.py 고정 규칙. 체크리스트·사업 프로필과 무관"],
        ],
    )

    add_heading(doc, "3. 목표 구조 (To-Be)", 1)
    add_heading(doc, "3.1 검토 단위: 「자료」→「사업(케이스)」", 2)
    add_para(doc, "[사업 검토 1건] (case_id)", bold=True)
    add_bullets(
        doc,
        [
            "사업 프로필 (보조금 유형, 비목 구성, 집행 특성 등)",
            "첨부: ①사업계획서 ②집행내역서 ③지출증빙 ④정산보고서",
            "통합 체크리스트 1개",
            "통합 결과 1건 (항목별 판정·근거·이력)",
        ],
    )

    add_heading(doc, "3.2 통합 체크리스트 + 적용 조건", 2)
    add_table(
        doc,
        ["속성", "역할"],
        [
            ["target_docs", "점검에 필요한 자료 (①만 / ②+③ / 4종 교차 등)"],
            ["applies_when", "적용 조건 (예: 인건비 비목 있음)"],
            ["exclude_when", "제외 조건 (예: 공사비만 집행)"],
            ["check_method", "교차대조 / 단일자료 규칙 / 유사도 보조 / 확인필요"],
            ["knowledge_ref", "지식DB·법령·감사사례 근거 링크"],
        ],
    )
    add_bullets(
        doc,
        [
            "공통 코어 + 모듈(증빙·정산 등) 구조 검토",
            "핵심 50~80개 + 사업유형별 서브셋이 현실적",
            "기존 data_type별 체크리스트·이력 마이그레이션·병행 기간 정책 필요",
        ],
    )

    add_heading(doc, "3.3 사업 프로필 (보완점 2의 핵심)", 2)
    add_para(doc, "체크리스트를 늘리기보다 적용 범위를 좁히고 N/A 처리가 우선입니다.")
    add_bullets(
        doc,
        [
            "보조금 유형: 시설보조 / 운영보조 / 행사보조 / 공사 중심 등",
            "집행 비목: 인건비, 용역비, 공사비, 포상금, 식비 등 (있음/없음)",
            "판정: P / F / W 외 N/A(해당 없음) / 확인필요",
        ],
    )

    add_heading(doc, "4. 권장 진행 순서 (4단계)", 1)
    add_para(
        doc,
        "원칙: 한 번에 전면 개편하지 않고, 오판 감소 → 통합 검토 → 지식·판정 고도화 순. "
        "2단계를 3단계보다 먼저 (통합만 하면 오판이 더 커질 수 있음).",
        bold=True,
    )

    phases = [
        (
            "1단계: 요구사항·도메인 모델 고정 (코드 전, 약 2~4주)",
            "개발 착수 전 업무·데이터 모델 합의",
            [
                "사업 검토 1건 정의 (4종 필수 여부, 미첨부 시 N/A)",
                "통합 체크리스트 분류 체계",
                "기존 ①~④ 항목 → 통합 항목 매핑표",
                "사업 프로필 항목 목록",
                "프로필–항목 매트릭스 (N/A 조건)",
                "판정 유형 정의 (P/F/W/N/A/확인필요)",
                "골든 케이스 2~3건 문서화",
            ],
            "완료: 담당자 검토 가능한 업무가이드 + 데이터 모델 초안 + 체크리스트 개편안",
        ),
        (
            "2단계: 사업 프로필 + 항목 적용 조건",
            "SILO 유지, 오판부터 감소 (리스크: 낮음)",
            [
                "검토 시작 시 사업 유형·비목 폼",
                "체크리스트 applies_when / exclude_when 메타데이터",
                "규칙검증 프로필 연동 (인건비·원천징수 등)",
                "N/A 항목 UI·종합 판정 규칙",
            ],
            "효과: 빠른 현장 신뢰도 회복. 3단계 전 필수 토대.",
        ),
        (
            "3단계: 사업(케이스) 단위 + 4종 통합",
            "SILO 해소 (리스크: 중~高)",
            [
                "case / business_review 엔티티",
                "4종 통합 업로드 UI",
                "유형별 파싱 유지 + 통합 비교 레이어",
                "통합 체크리스트 1개 (기존 4개 레거시)",
                "교차 검토 항목 (계획↔집행↔증빙↔정산)",
                "결과·이력 case_id 단일 관리",
            ],
            "현행 「전체 체크리스트 일괄」→ case 통합 검토로 대체.",
        ),
        (
            "4단계: 지식DB 연동 + 판정 고도화",
            "지식 기반 통합 체크리스트 완성 (리스크: 高)",
            [
                "항목 ↔ 지식DB 조항·감사사례 연결",
                "교차 판정 엔진 (수치·일자·금액 규칙 대조)",
                "정성 항목: 유사도 보조, 규칙·교차 중심",
                "Gemini: 검토 후 근거·조치안 (본문 미전송)",
            ],
            "",
        ),
    ]
    for title, goal, tasks, note in phases:
        add_heading(doc, title, 2)
        add_para(doc, f"목표: {goal}")
        add_bullets(doc, tasks)
        if note:
            add_para(doc, note)

    add_heading(doc, "교차 검토 항목 예시 (3단계)", 2)
    add_table(
        doc,
        ["점검 내용", "필요 자료"],
        [
            ["계획 예산 vs 집행 합계", "① + ②"],
            ["집행 항목 vs 증빙 금액·건수", "② + ③"],
            ["정산 총괄 vs 집행내역", "② + ④"],
            ["사업기간 vs 집행일·증빙일", "① + ② + ③"],
            ["자부담 비율 (계획 vs 정산)", "① + ④"],
        ],
    )

    add_heading(doc, "5. 보완점 ↔ 단계 매핑", 1)
    add_table(
        doc,
        ["보완점", "주 대응 단계"],
        [
            ["1. SILO", "3단계 (통합 case) + 4단계 (교차 판정)"],
            ["2. 사업 특성", "2단계 우선 (프로필·조건부 항목)"],
        ],
    )

    add_heading(doc, "6. 1단계 착수 전 체크리스트", 1)
    add_bullets(
        doc,
        [
            "□ 실제 검토 사업 2~3건 정리 (유형, 비목, 4종 유무, 오판 vs 정답)",
            "□ 교차 항목 10개 우선 선정",
            "□ 사업 프로필 질문 10문항 초안",
            "□ 통합 체크리스트 카테고리 합의",
            "□ 기존 DB 체크리스트·이력 마이그레이션 방침",
        ],
    )

    add_heading(doc, "7. 로드맵 한 줄 요약", 1)
    add_para(
        doc,
        "[1] 요구·골든케이스·통합 체크리스트·프로필 매트릭스 문서화\n"
        "      ↓\n"
        "[2] 사업 프로필 + 항목 적용조건 + N/A (오판 감소)\n"
        "      ↓\n"
        "[3] 사업 단위 4종 통합 검토 + 교차 항목 + 결과 통합 관리\n"
        "      ↓\n"
        "[4] 지식DB 항목 연동 + 교차/규칙 판정 고도화",
    )

    add_heading(doc, "8. 변경 이력", 1)
    add_table(
        doc,
        ["날짜", "내용"],
        [["2026-06-29", "초안 작성 — 검토 테스트 보완점 및 4단계 로드맵 정리"]],
    )

    add_heading(doc, "9. 메모 (추가 결정·회의록)", 1)
    add_para(doc, "(진행하면서 여기에 회의 결정·예외 사항을 적어 나가면 됩니다.)")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    saved = []
    for name in OUT_NAMES:
        path = OUT_DIR / name
        doc.save(path)
        saved.append(path)
    return saved[0]


if __name__ == "__main__":
    path = build()
    print(path)
